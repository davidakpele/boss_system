# app/routers/email_blast.py
"""
AI Email Campaign System
  GET  /email-campaigns               — Campaign dashboard
  GET  /email-campaigns/contacts      — Contact manager
  POST /email-campaigns/contacts/add  — Add single contact
  POST /email-campaigns/contacts/import — Bulk import CSV
  POST /email-campaigns/generate      — AI generates email from knowledge base
  GET  /email-campaigns/{id}          — Campaign detail / preview
  POST /email-campaigns/{id}/send     — Send now
  POST /email-campaigns/{id}/schedule — Schedule for later
  POST /email-campaigns/{id}/cancel   — Cancel scheduled
  GET  /email-campaigns/stats         — JSON stats
"""

import asyncio
import csv
import io
import logging
from datetime import datetime, timedelta

from fastapi import APIRouter, BackgroundTasks, Depends, Request, HTTPException, UploadFile, File, Form
from fastapi.responses import HTMLResponse, JSONResponse
from fastapi.templating import Jinja2Templates
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func

from app.database import get_db, AsyncSessionLocal
from app.models import (
    User, UserRole, EmailContact, EmailCampaign, EmailCampaignRecipient,
    KnowledgeChunk, AuditLog,
)
from app.auth import require_user
from app.services.ai_service import ai_service
from app.services.email_service import send_email
from app.config import settings

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/email-campaigns", tags=["email_campaigns"])
templates = Jinja2Templates(directory="app/templates")


# ══════════════════════════════════════════════════════════════════════════════
#  DASHBOARD
# ══════════════════════════════════════════════════════════════════════════════

@router.get("", response_class=HTMLResponse)
async def campaigns_dashboard(
    request: Request,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_user),
):
    if current_user.role not in (UserRole.super_admin, UserRole.admin, UserRole.manager):
        raise HTTPException(403)

    campaigns = (await db.execute(
        select(EmailCampaign).order_by(EmailCampaign.created_at.desc()).limit(30)
    )).scalars().all()

    total_contacts = (await db.execute(
        select(func.count(EmailContact.id)).where(EmailContact.is_active == True)
    )).scalar() or 0

    total_sent = (await db.execute(
        select(func.sum(EmailCampaign.sent_count))
    )).scalar() or 0

    return templates.TemplateResponse(
        request=request,
        name="email_campaigns/dashboard.html",
        context={
            "user": current_user,
            "page": "email_campaigns",
            "campaigns": campaigns,
            "total_contacts": total_contacts,
            "total_sent": total_sent,
            "smtp_enabled": settings.smtp_enabled,
        },
    )


# ══════════════════════════════════════════════════════════════════════════════
#  CONTACTS
# ══════════════════════════════════════════════════════════════════════════════

@router.get("/contacts", response_class=HTMLResponse)
async def contacts_page(
    request: Request,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_user),
):
    contacts = (await db.execute(
        select(EmailContact).order_by(EmailContact.name.asc())
    )).scalars().all()

    return templates.TemplateResponse(
        request=request,
        name="email_campaigns/contacts.html",
        context={"user": current_user, "page": "email_campaigns", "contacts": contacts},
    )


@router.post("/contacts/add")
async def add_contact(
    request: Request,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_user),
):
    body = await request.json()
    email = (body.get("email") or "").strip().lower()
    if not email:
        return JSONResponse({"error": "Email required"}, status_code=400)

    if "@" not in email or "." not in email.split("@")[-1]:
        return JSONResponse({"error": "Invalid email address"}, status_code=400)

    existing = (await db.execute(
        select(EmailContact).where(EmailContact.email == email)
    )).scalar_one_or_none()
    if existing:
        return JSONResponse({"error": "Contact already exists"}, status_code=409)

    name = (body.get("name") or "").strip() or email.split("@")[0]

    c = EmailContact(
        name        = name,
        email       = email,
        title       = (body.get("title") or "").strip(),
        department  = (body.get("department") or "").strip(),
        institution = (body.get("institution") or "").strip(),
        tags        = body.get("tags") or [],
        notes       = (body.get("notes") or "").strip(),
        created_by  = current_user.id,
    )
    db.add(c)
    await db.commit()
    await db.refresh(c)
    return JSONResponse({"status": "created", "id": c.id, "name": c.name, "email": c.email})

@router.post("/contacts/import")
async def import_contacts_csv(
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_user),
):
    import re
    import fitz  # pymupdf
    from docx import Document as DocxDocument

    content = await file.read()
    filename = (file.filename or "").lower()
    text = ""

    # ── Extract raw text based on file type ──
    if filename.endswith(".pdf"):
        try:
            pdf = fitz.open(stream=content, filetype="pdf")
            text = "\n".join(page.get_text() for page in pdf)
        except Exception as e:
            return JSONResponse({"error": f"Could not read PDF: {e}"}, status_code=400)

    elif filename.endswith(".docx"):
        try:
            from io import BytesIO
            doc = DocxDocument(BytesIO(content))
            text = "\n".join(p.text for p in doc.paragraphs)
            # also grab text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += "\n" + cell.text
        except Exception as e:
            return JSONResponse({"error": f"Could not read DOCX: {e}"}, status_code=400)

    elif filename.endswith(".csv") or filename.endswith(".txt"):
        text = content.decode("utf-8-sig", errors="replace")

    else:
        return JSONResponse(
            {"error": "Unsupported file type. Please upload CSV, TXT, PDF, or DOCX."},
            status_code=400,
        )

    # ── Extract all emails via regex ──
    emails_found = list(dict.fromkeys(  # deduplicate, preserve order
        m.lower() for m in re.findall(
            r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}",
            text
        )
    ))

    if not emails_found:
        return JSONResponse({"error": "No valid email addresses found in the file."}, status_code=400)

    created = 0
    skipped = 0

    for email in emails_found:
        existing = (await db.execute(
            select(EmailContact).where(EmailContact.email == email)
        )).scalar_one_or_none()
        if existing:
            skipped += 1
            continue
        db.add(EmailContact(
            name       = email.split("@")[0],  # fallback name
            email      = email,
            created_by = current_user.id,
        ))
        created += 1

    await db.commit()
    return JSONResponse({
        "created": created,
        "skipped": skipped,
        "found":   len(emails_found),
    })

@router.delete("/contacts/{contact_id}")
async def delete_contact(
    contact_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_user),
):
    c = (await db.execute(
        select(EmailContact).where(EmailContact.id == contact_id)
    )).scalar_one_or_none()
    if c:
        c.is_active = False
        await db.commit()
    return JSONResponse({"status": "deactivated"})


@router.get("/contacts/list")
async def list_contacts_json(
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_user),
):
    contacts = (await db.execute(
        select(EmailContact)
        .where(EmailContact.is_active == True)
        .order_by(EmailContact.name.asc())
    )).scalars().all()
    return JSONResponse([{
        "id":          c.id,
        "name":        c.name,
        "email":       c.email,
        "title":       c.title,
        "institution": c.institution,
        "department":  c.department,
        "tags":        c.tags or [],
    } for c in contacts])


# ══════════════════════════════════════════════════════════════════════════════
#  AI EMAIL GENERATION
# ══════════════════════════════════════════════════════════════════════════════

@router.post("/generate")
async def generate_email(
    request: Request,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_user),
):
    """
    AI generates a professional email using:
    - The user's prompt / purpose
    - Company knowledge base (RAG)
    - Tone, audience, and style preferences
    """
    body = await request.json()
    prompt      = (body.get("prompt") or "").strip()
    tone        = body.get("tone", "professional")       # professional|friendly|formal|persuasive
    audience    = body.get("audience", "")               # e.g. "Librarians in Nigerian universities"
    campaign_name = body.get("campaign_name", "")
    sender_name   = body.get("sender_name", "")
    sender_email  = body.get("sender_email", "")
    sender_phone  = body.get("sender_phone", "")

    if not prompt:
        return JSONResponse({"error": "Prompt is required"}, status_code=400)

    # RAG: pull relevant knowledge base context
    context_chunks = await ai_service.retrieve_context(prompt, db)
    context_text   = "\n\n".join(
        c.get("content", "")[:300] for c in context_chunks[:3]
    ) if context_chunks else ""

    # Build AI prompt
    system_prompt = f"""You are a professional business email writer for a company.
Your task is to write a complete, compelling email that represents the company well.

RULES:
- Write in a {tone} tone
- The email must be factual, professional, and persuasive where appropriate
- Use the company knowledge base context provided below as the source of truth
- Do NOT invent facts — only use what is in the context
- Format the email clearly: greeting, body paragraphs, call to action, sign-off
- Return ONLY the email body text, no subject line, no JSON wrapper
- Use **bold** for important terms (will render as HTML bold)
- The sign-off should use the sender details provided

COMPANY KNOWLEDGE BASE:
{context_text if context_text else "No specific context available — write based on the prompt."}

SENDER DETAILS:
Name: {sender_name}
Email: {sender_email}
Phone: {sender_phone}
"""

    user_message = f"""Write a professional email with the following purpose:

PURPOSE: {prompt}
TARGET AUDIENCE: {audience or "General business contacts"}
TONE: {tone}

Generate a complete email body ready to send."""

    messages = [
        {"role": "user", "content": f"{system_prompt}\n\n{user_message}"},
    ]

    try:
        email_body = await ai_service.chat_complete(messages)
        logger.info(f"AI raw response length: {len(email_body) if email_body else 0}")
        logger.info(f"AI raw response: {repr(email_body[:200]) if email_body else 'EMPTY'}")
        email_body = email_body.strip()
    except Exception as e:
        logger.error(f"AI email generation error: {e}", exc_info=True)
        return JSONResponse({"error": "AI failed to generate email. Check Ollama."}, status_code=500)

    if not email_body:
        return JSONResponse({"error": "AI returned empty response"}, status_code=500)

    subject_msgs = [
        {"role": "system", "content": "You write concise, compelling email subject lines. Return ONLY the subject line, nothing else."},
        {"role": "user",   "content": f"Write a subject line for this email purpose: {prompt}"},
    ]
    try:
        subject = (await ai_service.chat_complete(subject_msgs)).strip().strip('"').strip("'")
    except Exception:
        subject = f"Important Message from {sender_name or 'Our Company'}"
    import re
    html_body = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', email_body)
    html_body = html_body.replace('\n\n', '</p><p>').replace('\n', '<br>')
    html_body = f"<p>{html_body}</p>"

    return JSONResponse({
        "subject":          subject,
        "text_body":        email_body,
        "html_body":        html_body,
        "context_used":     bool(context_text),
        "knowledge_chunks": len(context_chunks) if context_chunks else 0,
    })


# ══════════════════════════════════════════════════════════════════════════════
#  SAVE CAMPAIGN
# ══════════════════════════════════════════════════════════════════════════════

@router.post("/save")
async def save_campaign(
    request: Request,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_user),
):
    body = await request.json()

    campaign = EmailCampaign(
        name           = body.get("campaign_name", "Untitled Campaign"),
        subject        = body.get("subject", ""),
        html_body      = body.get("html_body", ""),
        text_body      = body.get("text_body", ""),
        ai_prompt      = body.get("prompt", ""),
        status         = "draft",
        created_by     = current_user.id,
        total_recipients = 0,
    )
    db.add(campaign)
    await db.flush()

    # Save recipients
    recipient_ids = body.get("recipient_ids", [])
    recipient_emails = body.get("recipient_emails", [])  # manual emails

    count = 0
    for cid in recipient_ids:
        contact = (await db.execute(
            select(EmailContact).where(EmailContact.id == int(cid))
        )).scalar_one_or_none()
        if contact:
            db.add(EmailCampaignRecipient(
                campaign_id = campaign.id,
                contact_id  = contact.id,
                email       = contact.email,
                name        = contact.name,
            ))
            count += 1

    # Add any manually typed emails
    for em in recipient_emails:
        em = em.strip()
        if em and "@" in em:
            db.add(EmailCampaignRecipient(
                campaign_id = campaign.id,
                email       = em,
            ))
            count += 1

    campaign.total_recipients = count
    await db.commit()
    await db.refresh(campaign)

    return JSONResponse({"status": "saved", "campaign_id": campaign.id, "recipients": count})


# ══════════════════════════════════════════════════════════════════════════════
#  SEND NOW
# ══════════════════════════════════════════════════════════════════════════════

@router.post("/{campaign_id}/send")
async def send_campaign(
    campaign_id: int,
    background_tasks: BackgroundTasks,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_user),
):
    if not settings.smtp_enabled:
        return JSONResponse({"error": "SMTP not configured. Add SMTP_HOST, SMTP_USER, SMTP_PASSWORD to .env"}, status_code=503)

    campaign = (await db.execute(
        select(EmailCampaign).where(EmailCampaign.id == campaign_id)
    )).scalar_one_or_none()
    if not campaign:
        raise HTTPException(404)

    if campaign.status in ("sending", "sent"):
        return JSONResponse({"error": f"Campaign is already {campaign.status}"}, status_code=400)

    campaign.status     = "sending"
    campaign.started_at = datetime.utcnow()
    await db.commit()

    background_tasks.add_task(_send_campaign_emails, campaign_id)

    return JSONResponse({"status": "sending", "campaign_id": campaign_id})


async def _send_campaign_emails(campaign_id: int):
    """Background task: send emails to all campaign recipients."""
    async with AsyncSessionLocal() as db:
        campaign = (await db.execute(
            select(EmailCampaign).where(EmailCampaign.id == campaign_id)
        )).scalar_one_or_none()
        if not campaign:
            return

        recipients = (await db.execute(
            select(EmailCampaignRecipient)
            .where(
                EmailCampaignRecipient.campaign_id == campaign_id,
                EmailCampaignRecipient.status == "pending",
            )
        )).scalars().all()

        sent = 0
        failed = 0

        for r in recipients:
            # Personalise the email for each recipient
            personalised_html = campaign.html_body
            personalised_text = campaign.text_body or ""

            if r.name:
                # Replace generic greeting with personalised one
                personalised_html = personalised_html.replace(
                    "Dear Esteemed", f"Dear {r.name},"
                ).replace("Dear Sir/Madam", f"Dear {r.name}")
                personalised_text = personalised_text.replace(
                    "Dear Esteemed", f"Dear {r.name},"
                ).replace("Dear Sir/Madam", f"Dear {r.name}")

            success = await send_email(
                to_email  = r.email,
                to_name   = r.name or "",
                subject   = campaign.subject,
                html_body = personalised_html,
                text_body = personalised_text,
            )

            r.status  = "sent" if success else "failed"
            r.sent_at = datetime.utcnow() if success else None
            if success:
                sent += 1
            else:
                failed += 1
                r.error = "SMTP send failed"

            await db.commit()
            # Small delay between emails to avoid rate limiting
            await asyncio.sleep(0.3)

        campaign.sent_count   = sent
        campaign.failed_count = failed
        campaign.status       = "sent"
        campaign.completed_at = datetime.utcnow()
        await db.commit()

        logger.info(f"Campaign {campaign_id} complete: {sent} sent, {failed} failed")


# ══════════════════════════════════════════════════════════════════════════════
#  SCHEDULE CAMPAIGN
# ══════════════════════════════════════════════════════════════════════════════

@router.post("/{campaign_id}/schedule")
async def schedule_campaign(
    campaign_id: int,
    request: Request,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_user),
):
    body = await request.json()
    scheduled_at_str = body.get("scheduled_at", "")

    try:
        scheduled_at = datetime.fromisoformat(scheduled_at_str.replace("Z", "+00:00"))
        if scheduled_at.tzinfo:
            scheduled_at = scheduled_at.replace(tzinfo=None)
    except Exception:
        return JSONResponse({"error": "Invalid date format"}, status_code=400)

    if scheduled_at <= datetime.utcnow():
        return JSONResponse({"error": "Scheduled time must be in the future"}, status_code=400)

    campaign = (await db.execute(
        select(EmailCampaign).where(EmailCampaign.id == campaign_id)
    )).scalar_one_or_none()
    if not campaign:
        raise HTTPException(404)

    campaign.status       = "scheduled"
    campaign.scheduled_at = scheduled_at
    await db.commit()

    return JSONResponse({
        "status":       "scheduled",
        "scheduled_at": scheduled_at.isoformat(),
    })


@router.post("/{campaign_id}/cancel")
async def cancel_campaign(
    campaign_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_user),
):
    campaign = (await db.execute(
        select(EmailCampaign).where(EmailCampaign.id == campaign_id)
    )).scalar_one_or_none()
    if campaign and campaign.status == "scheduled":
        campaign.status = "draft"
        await db.commit()
    return JSONResponse({"status": "cancelled"})


@router.get("/{campaign_id}/detail")
async def campaign_detail(
    campaign_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_user),
):
    campaign = (await db.execute(
        select(EmailCampaign).where(EmailCampaign.id == campaign_id)
    )).scalar_one_or_none()
    if not campaign:
        raise HTTPException(404)

    recipients = (await db.execute(
        select(EmailCampaignRecipient)
        .where(EmailCampaignRecipient.campaign_id == campaign_id)
        .limit(100)
    )).scalars().all()

    return JSONResponse({
        "id":           campaign.id,
        "name":         campaign.name,
        "subject":      campaign.subject,
        "html_body":    campaign.html_body,
        "status":       campaign.status,
        "sent_count":   campaign.sent_count,
        "failed_count": campaign.failed_count,
        "total":        campaign.total_recipients,
        "scheduled_at": campaign.scheduled_at.isoformat() if campaign.scheduled_at else None,
        "created_at":   campaign.created_at.isoformat() if campaign.created_at else None,
        "recipients":   [{"email": r.email, "name": r.name, "status": r.status} for r in recipients],
    })