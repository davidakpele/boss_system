# src/app/services/document_service.py
import os
import csv
import io
from pathlib import Path
from typing import Optional
import logging

logger = logging.getLogger(__name__)


async def extract_text_from_file(file_path: str, file_type: str) -> str:
    """Extract text content from PDF, DOCX, or CSV files."""
    try:
        if file_type == "pdf":
            return await _extract_pdf(file_path)
        elif file_type in ("docx", "doc"):
            return await _extract_docx(file_path)
        elif file_type == "csv":
            return await _extract_csv(file_path)
        elif file_type == "txt":
            return await _extract_txt(file_path)
        else:
            return ""
    except Exception as e:
        logger.error(f"File extraction error for {file_path}: {e}")
        return ""


async def _extract_pdf(file_path: str) -> str:
    try:
        import PyPDF2
        text_parts = []
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        return "\n\n".join(text_parts)
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        return ""


async def _extract_docx(file_path: str) -> str:
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return "\n\n".join(paragraphs)
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        return ""


async def _extract_csv(file_path: str) -> str:
    try:
        import pandas as pd
        df = pd.read_csv(file_path, encoding="utf-8", errors="replace")
        return df.to_string(index=False)
    except Exception as e:
        logger.error(f"CSV extraction error: {e}")
        return ""


async def _extract_txt(file_path: str) -> str:
    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()
    except Exception as e:
        logger.error(f"TXT extraction error: {e}")
        return ""


def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> list[str]:
    """Split text into overlapping chunks."""
    if not text:
        return []
    words = text.split()
    chunks = []
    i = 0
    while i < len(words):
        chunk_words = words[i:i + chunk_size]
        chunks.append(" ".join(chunk_words))
        i += chunk_size - overlap
    return chunks


def get_file_type(filename: str) -> str:
    ext = Path(filename).suffix.lower().lstrip(".")
    mapping = {
        "pdf": "pdf",
        "docx": "docx",
        "doc": "docx",
        "csv": "csv",
        "txt": "txt",
        "xls": "excel",
        "xlsx": "excel",
    }
    return mapping.get(ext, "unknown")